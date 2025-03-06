import streamlit as st
import re
import pandas as pd
from docx import Document
import nltk
from nltk.tokenize import sent_tokenize
from nltk.corpus import stopwords
import torch
import torch.nn.functional as F
from transformers import AutoModelForSequenceClassification, AutoTokenizer
from transformers import BertTokenizer, BertForSequenceClassification
import spacy


st.set_page_config(
    layout="wide",
    page_title="Conference Calls Transcript Sentiment Analyzer",
    page_icon="📊",
)

nltk.download("punkt")
nltk.download("punkt_tab")
nltk.download("stopwords")


@st.cache_resource
def load_sentiment_model():
    sentiment_model_name = "ProsusAI/finbert"
    sentiment_tokenizer = BertTokenizer.from_pretrained(sentiment_model_name)
    sentiment_model = BertForSequenceClassification.from_pretrained(
        sentiment_model_name
    )
    sentiment_model.eval()
    return sentiment_tokenizer, sentiment_model


sentiments = ["Positive", "Negative", "Neutral"]


@st.cache_resource
def load_emotion_model():
    emotion_model_name = "j-hartmann/emotion-english-distilroberta-base"
    emotion_tokenizer = AutoTokenizer.from_pretrained(emotion_model_name)
    emotion_model = AutoModelForSequenceClassification.from_pretrained(
        emotion_model_name
    )
    emotion_model.eval()
    return emotion_tokenizer, emotion_model


emotions = ["anger", "disgust", "fear", "joy", "neutral", "sadness", "surprise"]


relevant_entities = [
    "ORG",
    "MONEY",
    "PERCENT",
    "CARDINAL",
    "GPE",
    "DATE",
    "EVENT",
    "PRODUCT",
    "LAW",
]


sentiment_tokenizer, sentiment_model = load_sentiment_model()
emotion_tokenizer, emotion_model = load_emotion_model()
nlp = spacy.load("en_core_web_md")


stop_words = stopwords.words("english")


def remove_stopwords(sentence, stopwords=None):
    pattern = r"\b(?:" + "|".join(re.escape(word) for word in stopwords) + r")\b"

    cleaned_sentence = re.sub(pattern, "", sentence, flags=re.IGNORECASE)

    return " ".join(cleaned_sentence.split())


def replace_apostrophe(text):
    text = re.sub("’", "'", text)
    return text


custom_stopwords = set(stopwords.words("english")) | set(
    [
        "um",
        "uh",
        "okay",
        "well",
        "you know",
        "i mean",
        "like",
        "think",
        "question",
        "questions",
        "so",
        "actually",
        "basically",
        "just",
        "right",
        "sure",
        "yeah",
        "yep",
        "nope",
        "great",
        "thanks",
        "thank you",
        "good morning",
        "good afternoon",
        "good evening",
        "ladies and gentlemen",
        "everyone",
        "folks",
        "team",
        "welcome",
        "pleasure",
        "appreciate",
        "introduction",
        "moving on",
        "next slide",
        "let’s move on",
        "let’s begin",
        "let’s get started",
        "first of all",
        "secondly",
        "last but not least",
        "before we start",
        "before we begin",
        "without further ado",
        "i’d like to",
        "we’d like to",
        "going forward",
        "as you can see",
        "as mentioned earlier",
        "as we discussed",
        "as previously stated",
        "as i said",
        "as we said",
        "again",
        "also",
        "furthermore",
        "moreover",
        "in addition",
        "of course",
        "obviously",
        "clearly",
        "frankly",
        "honestly",
        "to be honest",
        "to be frank",
        "to be clear",
        "as far as i know",
        "as far as we know",
        "our perspective",
        "our standpoint",
        "to some extent",
        "at the end of the day",
        "bottom line",
        "high level",
        "big picture",
        "to wrap up",
        "to summarize",
        "in conclusion",
        "in summary",
        "before i hand it over",
        "before we wrap up",
        "any questions",
        "do you have any questions",
        "q&a",
        "let’s take questions",
        "moving to the next question",
        "let me check",
        "we will get back to you",
        "we will follow up",
        "circle back",
        "touch base",
        "take this offline",
        "offline discussion",
        "follow-up",
        "let’s revisit",
        "we are looking into it",
        "we are working on it",
        "stay tuned",
        "more details to come",
        "we don’t have that information right now",
        "i don’t have that data",
        "good question",
        "great question",
        "let me clarify",
        "let me add",
        "if i may",
        "if you will",
        "if you look at",
        "when you think about it",
        "at this point in time",
        "right now",
        "currently",
        "as of now",
        "in terms of",
        "with respect to",
        "regarding",
        "pertaining to",
        "relative to",
        "looking at",
        "focusing on",
        "from a standpoint of",
        "from a perspective of",
        "from an angle of",
        "talking about",
        "discussing",
        "speaking of",
        "with that said",
        "having said that",
        "on that note",
        "before i forget",
        "one more thing",
        "please",
        "operator",
    ]
)


def clean_text(text):
    text = text.lower()
    text = re.sub(r"\bapple intelligence\b", "appleintelligence", text)
    text = re.sub(r"\byear-over-year\b", "yearoveryear", text)
    text = re.sub(r"\ball-time\b", "alltime", text)
    text = re.sub(r"\ball-in-one\b", "allinone", text)
    text = re.sub(r"(?<!\w)(u\.s\.)(?!\w)", "unitedstates", text)
    text = re.sub(r"(?<!\w)(u\.k\.)(?!\w)", "unitedkingdom", text)
    text = re.sub(r"(?<!\w)(i\.e\.)(?!\w)", "that is", text)
    text = re.sub(r"\biphone (1[0-6]|[1-9])\b", lambda m: f"iphone{m.group(1)}", text)
    text = replace_apostrophe(text)
    text = remove_stopwords(text, stopwords=custom_stopwords)
    text = re.sub(r"[^\w\s%$.]", "", text)
    text = re.sub(r"\.$", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def extract_name_role(pairs):
    name_role_dict = {}
    for pair in pairs:
        name, role = [item.strip() for item in pair.split(" - ")]
        name_role_dict[name] = role
    return name_role_dict


def extract_dialogues(doc, names):
    dialogues = []
    current_speaker = None
    current_dialogue = []

    for para in doc.paragraphs:
        line = para.text.strip()

        if line in names:
            if current_speaker:
                dialogues.append((current_speaker, " ".join(current_dialogue)))
            current_speaker = line
            current_dialogue = []
        else:
            current_dialogue.append(line)

    if current_speaker and current_dialogue:
        dialogues.append((current_speaker, " ".join(current_dialogue)))

    return dialogues


def get_sentiments(list_of_dialogues):
    inputs = sentiment_tokenizer(
        list_of_dialogues, return_tensors="pt", padding=True, truncation=True
    )

    with torch.no_grad():
        output = sentiment_model(**inputs)

    logits = output.logits
    probabilities = torch.nn.functional.softmax(logits, dim=-1)

    predicted_indices = torch.argmax(probabilities, dim=1)
    sentiment_labels = [sentiments[idx] for idx in predicted_indices.tolist()]

    confidence = probabilities[
        torch.arange(len(predicted_indices)), predicted_indices
    ].tolist()

    sentence_scores = (probabilities[:, 0] - probabilities[:, 1]).tolist()

    return (
        sentiment_labels,
        confidence,
        sentence_scores,
    )


def classify_emotions(statement):
    inputs = emotion_tokenizer(
        statement, return_tensors="pt", truncation=True, padding=True
    )
    with torch.no_grad():
        outputs = emotion_model(**inputs)
        probs = F.softmax(outputs.logits, dim=-1)
        predicted_label = emotions[torch.argmax(probs).item()]
        return predicted_label


def remove_stop_words(text):
    text = remove_stopwords(text, stopwords=stop_words)
    return text.strip()


def get_strategic_focus_areas(text):
    doc = nlp(text)

    entities = []

    for ent in doc.ents:
        if ent.label_ in relevant_entities:
            entities.append(remove_stop_words(ent.text))

    if entities:
        return ", ".join(entities)
    else:
        return "No Strategic Focus Areas"


def color_sentiment(val):
    if val == "Positive":
        return "background-color: lightgreen; color: black"
    elif val == "Neutral":
        return "background-color: yellow; color: black"
    elif val == "Negative":
        return "background-color: red; color: white"
    return ""


st.sidebar.markdown(
    """
        <style>
            .main-title {
                text-align: center;
                font-size: 1.5em;
                color: #4A90E2;
                font-weight: bold;
            }

            .sub-heading {
                text-align: center;
                font-size: 1.2em;
                color: grey;
            }

            .info-card {
                background-color: #F5F5F5;
                padding: 20px;
                border-radius: 20px;
                box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
                margin: 20px 0;
            }

            .feature-list {
                font-size: 1.2em;
                color: #2C3E50;
                line-height: 1.8;
            }

        </style>
    """,
    unsafe_allow_html=True,
)

st.sidebar.markdown(
    '<div class="main-title">Conference Calls Transcript Sentiment Analyzer 📊</div>',
    unsafe_allow_html=True,
)

st.sidebar.markdown(
    '<div class="sub-heading">Gain Valuable Insights from Conference Call Transcripts</div>',
    unsafe_allow_html=True,
)

st.sidebar.markdown(
    """
    <div class="info-card">
        <p class="feature-list">📤 <strong>Upload & Analyze:</strong> Easily upload transcripts and receive <em>instant sentiment analysis</em>.</p>
        <p class="feature-list">📊 <strong>Sentiment Detection:</strong> Identify whether the sentiment is <em>Positive</em>, <em>Negative</em>, or <em>Neutral</em>.</p>
        <p class="feature-list">😀 <strong>Emotion Detection:</strong> Classify the statement into one of the following emotions; <em>Anger</em>, <em>Disgust</em>, <em>Fear</em>, <em>Joy</em>, <em>Neutral</em>, <em>Sadness</em>, or <em>Surprise</em>.</p>
        <p class="feature-list">📈 <strong>Interactive Visualization:</strong> Explore sentiment trends through <em>dynamic charts</em> and <em>graphs</em>.</p>
    </div>
    """,
    unsafe_allow_html=True,
)

uploaded_file = st.sidebar.file_uploader("Upload a Word Document (DOCX)", type=["docx"])

st.sidebar.markdown(
    f"<p style='color: grey;'>Please refresh the page before uploading a new Word document to ensure that any cached data from the previously uploaded document is cleared.</p>",
    unsafe_allow_html=True,
)

st.subheader("Uploaded File Content")

if uploaded_file is not None:
    doc = Document(uploaded_file)

    company_participants = []
    company_participants_index = -1

    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text == "Company Participants":
            company_participants_index = i
        if i == (company_participants_index + 1) and company_participants_index != -1:
            company_participants.append(paragraph.text.split("\n"))

    analysts = []
    conference_call_participants_index = -1

    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text == "Conference Call Participants":
            conference_call_participants_index = i
        if (
            i == (conference_call_participants_index + 1)
            and conference_call_participants_index != -1
        ):
            analysts.append(paragraph.text.split("\n"))

    company_participants_dictionary = (
        extract_name_role(company_participants[0]) if company_participants else {}
    )
    company_participants_names = list(company_participants_dictionary.keys())
    analyst_names = (
        [name.split("-")[0].strip() for name in analysts[0]] if analysts else []
    )

    names = company_participants_names + analyst_names + ["Operator"]

    result = extract_dialogues(doc, names)

    df = pd.DataFrame(result, columns=["Speaker", "Statement"])

    def map_role(name):
        if name in company_participants_dictionary:
            return company_participants_dictionary[name]
        elif name in analyst_names:
            return "Analyst"
        else:
            return "Operator"

    df["Role"] = df["Speaker"].apply(map_role)

    rows = []

    for idx, row in df.iterrows():
        sentences = sent_tokenize(row["Statement"])
        for sentence in sentences:
            rows.append(
                {"Speaker": row["Speaker"], "Role": row["Role"], "Statement": sentence}
            )

    df = pd.DataFrame(rows)

    df["Cleaned Statement"] = df["Statement"].apply(clean_text)

    df = df[df["Cleaned Statement"].str.split().apply(len) >= 3].reset_index(drop=True)

    df.drop("Cleaned Statement", axis=1, inplace=True)

    st.session_state.df = df

if "df" in st.session_state and st.session_state.df is not None:
    st.dataframe(st.session_state.df)
else:
    st.markdown(
        f"<p style='color: grey;'>Please upload a valid Word document (DOCX) to view its content and perform sentiment analysis.</p>",
        unsafe_allow_html=True,
    )

if "df" in st.session_state and not st.session_state.df.empty and st.button("Analyze"):
    with st.spinner("Analyzing, please wait..."):
        (
            st.session_state.df["Sentiment Label"],
            st.session_state.df["Confidence Rate"],
            st.session_state.df["Sentiment Score"],
        ) = get_sentiments(st.session_state.df["Statement"].tolist())
        st.session_state.df["Emotion"] = st.session_state.df["Statement"].apply(
            classify_emotions
        )
        st.session_state.df["Strategic Focus Areas"] = st.session_state.df[
            "Statement"
        ].apply(get_strategic_focus_areas)

    st.session_state.labeled_df = st.session_state.df

if "labeled_df" in st.session_state and st.session_state.labeled_df is not None:
    st.divider()
    st.subheader("Analyzed Content")
    st.dataframe(
        st.session_state.labeled_df.style.map(
            color_sentiment, subset=["Sentiment Label"]
        )
    )

    st.markdown(
        f"<p style='color: grey;'>Download the table above and upload it to Power BI to generate a sentiment analysis report with visualizations.</p>",
        unsafe_allow_html=True,
    )

    csv = st.session_state.labeled_df.to_csv(index=False).encode("utf-8")

    st.download_button(
        label="Download CSV",
        data=csv,
        file_name="conference_call_trancripts.csv",
        mime="text/csv",
    )

st.sidebar.markdown(
    f"<p style='color: grey; text-align: center;'>made by <a style = 'text-decoration: None;' target='_blank' href= 'https://github.com/razanaqvi14'>Ali Raza</a></p>",
    unsafe_allow_html=True,
)
